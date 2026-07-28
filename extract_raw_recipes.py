#!/usr/bin/env python3
"""
从 私房菜谱.docx 提取所有菜谱标题，生成 未拆分菜谱清单.md
"""
import re
from pathlib import Path
import docx

SKIP = {
    '私房菜谱', '无肉不欢', '碳水の诱惑', '汤的诱惑', '菜菜子',
    '凉菜', '厨房小技巧', '火锅', '炸货', '附录',
    '📝 总结与通用建议', '24款万能酱汁', '厨房酱料完整汇总表',
    '六款万能蘸水详细步骤表', '万能水煮菜清单（适配以上所有蘸水）',
    '常用淀粉与面粉', '肉馅对比', '厨房常见酱料作用', '辣椒油制作',
    '食材处理技巧', '烹饪技巧', '调味技巧', '其他技巧', '刷蜂蜜', '裹糊',
    '料酒和黄酒的区别', '东北酸甜酱', '酸奶薄荷酱', '万能凉拌汁',
    '东北烤肉-酸甜水料', '日式酱油', '家庭版味淋', '酸辣汁',
    '泰式风味烤肉酱料汇总', '粤菜/港式风味烤肉酱料汇总', '卤牛肉蘸汁',
    '拌面酱料', '《6款无敌蘸水：点亮水煮菜的灵魂》',
    '解腻盐渍水果：万能公式+经典做法大全',
    '凉拌菜核心心法：万能公式与基础调味汁',
    '户外木炭烧烤全攻略', '家庭烤肉', '家庭电烤炉烤串', '吊炉烤肉',
    '【附】武大郎烧饼（山东阳谷·煎烙版）差异速记', '炸货卷饼',
    '菜花合集', '地三鲜', '地三鲜–省油版',
}


def clean_text(text):
    """去除飞书导出导致的重复文本（3次重复）"""
    t = text.strip()
    if not t:
        return ''
    for n in [3, 2]:
        part = len(t) // n
        if part > 4 and t[:part] == t[part:2*part]:
            if n == 2 or (n == 3 and t[:part] == t[2*part:3*part]):
                return t[:part].strip()
    return t


def extract_raw_recipes(filepath):
    doc = docx.Document(filepath)
    body = list(doc.element.body)

    # 收集菜谱标题
    recipes = []
    current_title = None
    para_idx = 0

    for el in body:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag

        if tag == 'p':
            if para_idx >= len(doc.paragraphs):
                para_idx += 1
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1
            text = clean_text(para.text)
            if not text:
                continue

            # 获取字号
            sz = None
            for run in para.runs:
                if run.font.size:
                    sz = run.font.size
                    break

            # 菜谱标题（字号 >= 190000）
            if sz and sz >= 190000 and text not in SKIP and len(text) >= 2:
                if not re.match(r'^[0-9\s\.、，。！？]+$', text):
                    recipes.append({'title': text})

    print(f"提取到 {len(recipes)} 个菜谱")

    # 生成汇总表格：未拆分菜谱清单.md
    base_dir = Path(__file__).parent
    summary_path = base_dir / '未拆分菜谱清单.md'
    print(f"生成汇总表格: {summary_path}")

    summary_lines = [
        '# 未拆分菜谱清单',
        '',
        '此清单列出从 `私房菜谱.docx` 提取出的所有菜谱，拆分完成后在状态列标注 ✅。',
        '',
        '| 序号 | 菜谱标题 | 拆分状态 |',
        '| ---- | -------- | -------- |',
    ]

    for idx, recipe in enumerate(recipes, 1):
        title = recipe['title']
        summary_lines.append(f'| {idx} | {title} |  |')

    summary_content = '\n'.join(summary_lines) + '\n'
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(summary_content)
    print(f"✅ 已生成未拆分菜谱清单: {summary_path}")

    return len(recipes)


def main():
    base_dir = Path(__file__).parent
    docx_path = base_dir / '私房菜谱.docx'

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        return

    print(f"📖 读取: {docx_path}")
    count = extract_raw_recipes(str(docx_path))
    print(f"✅ 完成！共 {count} 个菜谱")


if __name__ == '__main__':
    main()