#!/usr/bin/env python3
"""
私房菜谱 .docx → recipes.json 解析脚本 v2
===========================================
处理飞书导出的特殊格式：段落内容在前，所有表格在文档末尾。
自动匹配表格到对应菜谱。
"""

import sys, json, re, argparse, os
import docx

SKIP = {
    '私房菜谱', '无肉不欢', '碳水の诱惑', '汤的诱惑', '菜菜子',
    '凉菜', '厨房小技巧', '火锅', '炸货', '附录',
    '📝 总结与通用建议', '24款万能酱汁', '厨房酱料完整汇总表',
    '六款万能蘸水详细步骤表', '万能水煮菜清单（适配以上所有蘸水）',
    '常用淀粉与面粉', '肉馅对比', '厨房常见酱料作用', '辣椒油制作',
    '食材处理技巧', '烹饪技巧', '调味技巧', '其他技巧', '刷蜂蜜', '裹糊',
    '料酒和黄酒的区别', '东北酸甜酱', '酸奶薄荷酱', '万能凉拌汁',
    '东北烤肉-酸甜水料', '日式酱油', '家庭版味淋', '酸辣汁',
    '泰式风味烤肉酱料汇总', '粤菜/港式风味烤肉酱料汇总', '卤牛肉蘸汁',
    '拌面酱料', '《6款无敌蘸水：点亮水煮菜的灵魂》',
    '解腻盐渍水果：万能公式+经典做法大全',
    '凉拌菜核心心法：万能公式与基础调味汁',
    '户外木炭烧烤全攻略', '家庭烤肉', '家庭电烤炉烤串', '吊炉烤肉',
    '【附】武大郎烧饼（山东阳谷·煎烙版）差异速记', '炸货卷饼',
    '菜花合集', '地三鲜', '地三鲜–省油版',
}


def clean_text(text):
    """去除飞书导出导致的重复文本（3次重复）"""
    t = text.strip()
    if not t:
        return ''
    for n in [3, 2]:
        part = len(t) // n
        if part > 4 and t[:part] == t[part:2*part]:
            if n == 2 or (n == 3 and t[:part] == t[2*part:3*part]):
                return t[:part].strip()
    return t


def extract_table_data(table):
    """提取表格为二维数组，清理零宽字符"""
    rows = []
    for row in table.rows:
        cells = [clean_text(c.text).replace('\u200b', '') for c in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def parse_ingredients(table_data):
    """从表格中提取食材"""
    if len(table_data) < 2:
        return []
    hdr = [h.lower() for h in table_data[0]]
    name_col = amount_col = 0
    note_col = -1
    for i, h in enumerate(hdr):
        if any(k in h for k in ['食材', '材料', '主料', '名称', '类别', '类型']):
            name_col = i
        elif any(k in h for k in ['用量', '分量', '建议用量', '数量']):
            amount_col = i
        elif any(k in h for k in ['作用', '说明', '处理', '要点', '方式', '解析']):
            note_col = i
    ings = []
    skip_words = ['点击', '查看', '图片', '方法', '步骤', '阶段', '关键', '问题', '注意']
    for row in table_data[1:]:
        name = row[name_col] if name_col < len(row) else ''
        if not name or any(k in name for k in skip_words):
            continue
        amount = row[amount_col] if amount_col < len(row) else ''
        note = row[note_col] if 0 <= note_col < len(row) else ''
        ings.append({'name': name, 'amount': amount, 'note': note})
    return ings


def extract_servings(text):
    """提取份量"""
    m = re.search(r'(\d+[-~]\d+\s*人份)', text) or re.search(r'(\d+\s*人份)', text)
    return m.group(1) if m else '2-3人份'


def extract_cooking_time(text):
    """从文本中提取烹饪时间（分钟）"""
    patterns = [
        (r'慢[炖煮焖]*\s*(\d+)[-~](\d+)\s*小时', lambda m: int(m.group(2)) * 60),
        (r'[炖焖煮烧卤]*\s*(\d+)[-~](\d+)\s*小时', lambda m: int(m.group(2)) * 60),
        (r'(\d+)[-~](\d+)\s*小时', lambda m: int(m.group(2)) * 60),
        (r'至少需要\s*(\d+)\s*小时', lambda m: int(m.group(1)) * 60),
        (r'(\d+)\s*小时', lambda m: int(m.group(1)) * 60),
        (r'慢[炖煮焖]*\s*(\d+)[-~]?(\d*)\s*分钟', lambda m: int(m.group(2) or m.group(1))),
        (r'[炖焖煮烧]\s*(\d+)\s*分钟', lambda m: int(m.group(1))),
        (r'(\d+)\s*分钟', lambda m: int(m.group(1))),
    ]
    for pat, fn in patterns:
        m = re.search(pat, text)
        if m:
            return fn(m)
    return 30


def extract_steps_time(text):
    """从步骤文本中提取烹饪时间"""
    times = []
    for pat in [
        r'慢炖\s*(\d+)[-~]?(\d*)\s*(小时|分钟)',
        r'[炖煮焖烧卤]\s*(\d+)[-~]?(\d*)\s*(小时|分钟)',
        r'至少.*?(\d+)\s*(小时|分钟)',
        r'(\d+)[-~](\d+)\s*(小时|分钟)',
        r'(\d+)\s*(小时|分钟)',
    ]:
        for m in re.finditer(pat, text):
            groups = m.groups()
            unit = groups[-1]  # 最后一个捕获组是单位
            # 找到数值部分
            nums = [int(g) for g in groups[:-1] if g]
            val = max(nums) if nums else 0
            if val > 0:
                if unit == '小时':
                    times.append(val * 60)
                elif val >= 5:
                    times.append(val)
    return max(times) if times else 0


def guess_subcategory(title):
    """根据标题推断子分类"""
    maps = {
        '猪肉': ['猪肉', '排骨', '五花肉', '肉末', '猪蹄', '猪脚', '猪耳', '梅花肉', '里脊',
                '猪排', '肘子', '红烧肉', '卤肉', '把子肉', '回锅肉', '锅包肉', '粉蒸肉',
                '扣肉', '脆皮', '叉烧', '肉丝', '蹄筋', '蒜泥白肉', '猪头', '脆哨', '小酥肉', '糖醋里脊'],
        '鸡肉': ['鸡', '宫保', '鸡丁', '鸡腿', '鸡翅', '鸡胸', '鸡爪', '鸡煲', '鸡排', '鸡块',
                '烧鸟', '鸡丝', '鸡饭', '鸡翼', '鸡公煲', '口水鸡', '椒麻鸡', '手撕鸡', '辣子鸡',
                '咸蛋黄鸡', '鸡脚', '茶香鸡', '咖喱鸡', '照烧鸡', '柠檬鸡', '虫草花蒸鸡', '黄焖鸡'],
        '牛肉': ['牛肉', '牛排', '牛腩', '牛腱', '牛板筋', '牛肋', '牛柳', '牛丼', '牛肉干',
                '牛肚', '牛舌', '牛', '土豆炖牛腩', '番茄牛腩'],
        '羊肉': ['羊肉', '羊排', '羊蝎子', '羊蹄', '羊腿', '羊', '手抓羊肉', '手抓饭', '孜然羊肉'],
        '海鲜': ['鱼', '虾', '蟹', '贝', '海鲜', '鲈鱼', '黄花鱼', '鱿鱼', '生蚝', '扇贝', '蛤蜊',
                '海肠', '龙虾', '螺', '鳗鱼', '蚵仔', '虾仁', '虾滑', '虾米', '小龙虾', '螺片',
                '章鱼', '墨鱼', '鲍鱼', '海参', '蛏子', '花甲', '海螺', '带鱼', '龙利鱼', '三文鱼'],
        '豆腐': ['豆腐', '豆花', '豆皮', '豆干', '豆豉', '腐竹', '豆芽', '豆角'],
        '蛋类': ['蛋', '炒蛋', '蒸蛋', '皮蛋', '茶叶蛋', '溏心蛋', '鸡蛋糕', '鸡蛋蒜', '蛋炒饭',
                '西红柿炒鸡蛋', '韭菜炒鸡蛋', '西葫芦鸡蛋饼', '骨汤冲蛋', '紫菜蛋花汤', '糖拌番茄'],
        '素菜': ['土豆', '茄子', '番茄', '西红柿', '白菜', '青菜', '辣椒', '四季豆', '包菜',
                '花菜', '蘑菇', '洋葱', '蒜苔', '芹菜', '油菜', '油麦菜', '地瓜', '玉米', '豇豆',
                '黄瓜', '蒜', '藕', '萝卜', '冬瓜', '南瓜', '笋', '木耳', '金针菇', '地三鲜',
                '炒合菜', '干煸四季豆', '醋溜豆芽', '酸辣土豆丝', '红烧茄子', '木须肉', '拔丝',
                '甘蓝', '擂椒皮蛋', '甘梅', '老醋花生', '娃娃菜', '大头菜', '蒜泥拆骨肉',
                '孜然洋葱', '青椒酿肉', '线椒鸡蛋', '土豆丝', '炝大头菜', '炒芹菜', '炒蒜苔',
                '炒洋葱', '蘑菇炒肉', '虾米炒油菜', '青椒炒豆皮', '香辣孜然豆皮', '麻汁油麦菜',
                '凉拌包菜', '凉拌娃娃菜', '蒜苔腊肉', '蒜苔', '手锤茄子', '花生'],
        '面食': ['面', '面条', '饺子', '饼', '包子', '馒头', '馄饨', '火烧', '锅贴', '烧麦',
                '河粉', '焖面', '卷饼', '披萨', '水煎包', '拌面', '炒面', '凉面', '炒饼', '小笼包',
                '肉龙', '掉渣饼', '葱油拌面', '炝锅面', '沙茶面', '红烧牛肉面', '大肉面', '酱香饼',
                '炒河粉', '醋卤面', '馒头片', '方便面', '烤包子', '鸡蛋灌饼', '油条', '麻酱凉面',
                '鸡丝凉面', '淀粉肠', '肉火烧', '肉饼', '肉夹馍', '灌汤包', '三不沾', '鸡蛋糕',
                '烧饼', '煎饼', '烙饼', '馅饼', '春饼', '煎包', '生煎', '锅盔', '馍', '馕', '酥饼'],
        '米饭': ['米饭', '炒饭', '盖饭', '煲仔饭', '拌饭', '粥', '焖饭', '手抓饭', '卤肉饭',
                '牛丼', '饭包', '烤肉拌饭', '咖喱鸡排饭', '咖喱土豆饭', '海南鸡饭', '隆江猪脚饭',
                '猪油拌饭', '溏心蛋海苔碎', '脆哨炒饭', '芒果糯米饭', '米村拌饭', '鸡蛋糕拌饭',
                '糁汤', '美龄粥', '南瓜粥', '红豆粥', '疙瘩汤', '海鲜疙瘩汤', '蟹虾粥', '大酱汤',
                '奶油蘑菇汤', '辣条黄金炒饭', '酱油炒饭', '香辣鱿鱼拌饭', '海鲜粉丝煲'],
    }
    for cat, kws in maps.items():
        if any(kw in title for kw in kws):
            return cat
    return '猪肉'


def extract_tags(title, ingredients):
    tags = set()
    kw_map = {
        '猪肉': ['猪肉', '排骨', '五花肉', '肉末', '猪蹄', '猪脚', '猪耳', '梅花肉', '里脊', '猪排', '肘子', '红烧肉', '卤肉', '把子肉', '回锅肉', '锅包肉', '粉蒸肉', '扣肉', '脆皮', '叉烧', '肉丝', '蹄筋', '蒜泥白肉', '猪头', '脆哨', '小酥肉', '糖醋里脊'],
        '鸡肉': ['鸡', '宫保', '鸡丁', '鸡腿', '鸡翅', '鸡胸', '鸡爪', '鸡煲', '鸡排', '鸡块', '烧鸟', '鸡丝', '鸡饭', '鸡翼', '鸡公煲', '口水鸡', '椒麻鸡', '手撕鸡', '辣子鸡', '咸蛋黄鸡'],
        '牛肉': ['牛肉', '牛排', '牛腩', '牛腱', '牛板筋', '牛肋', '牛柳', '牛丼', '牛肉干', '牛肚'],
        '羊肉': ['羊肉', '羊排', '羊蝎子', '羊蹄', '羊腿', '羊'],
        '海鲜': ['鱼', '虾', '蟹', '贝', '海鲜', '鲈鱼', '黄花鱼', '鱿鱼', '生蚝', '扇贝', '蛤蜊', '海肠', '龙虾', '螺', '鳗鱼', '蚵仔', '小龙虾'],
        '豆腐': ['豆腐', '豆花', '豆皮', '豆干', '豆豉', '腐竹', '豆芽', '豆角'],
        '蛋类': ['蛋', '炒蛋', '蒸蛋', '皮蛋', '茶叶蛋', '溏心蛋', '鸡蛋糕', '鸡蛋蒜'],
        '素菜': ['土豆', '茄子', '番茄', '西红柿', '白菜', '青菜', '辣椒', '四季豆', '包菜', '花菜', '蘑菇', '洋葱', '蒜苔', '芹菜', '油菜', '油麦菜', '地瓜', '玉米', '豇豆', '黄瓜', '蒜', '藕', '萝卜', '冬瓜', '南瓜', '笋', '木耳', '金针菇'],
        '面食': ['面', '面条', '饺子', '饼', '包子', '馒头', '馄饨', '火烧', '锅贴', '烧麦', '河粉', '焖面', '卷饼', '披萨', '水煎包', '拌面', '炒面', '凉面', '炒饼', '小笼包', '肉龙', '掉渣饼'],
        '米饭': ['米饭', '炒饭', '盖饭', '煲仔饭', '拌饭', '粥', '焖饭', '手抓饭', '卤肉饭', '牛丼'],
        '红烧': ['红烧', '卤', '红焖', '焖'],
        '清蒸': ['清蒸', '蒸', '桑拿'],
        '爆炒': ['爆炒', '快炒', '干煸', '生炒', '炒'],
        '炖煮': ['炖', '煮', '慢炖', '煲', '焖煮', '清炖'],
        '油炸': ['炸', '酥', '天妇罗', '油焖'],
        '凉拌': ['凉拌', '拌', '口水', '手撕', '蘸水'],
        '烧烤': ['烤', '烧烤', '炭烤', '吊炉'],
        '火锅': ['火锅', '麻辣烫', '冒菜', '毛血旺', '关东煮', '寿喜锅', '锅'],
        '麻辣': ['麻辣', '花椒', '麻椒', '辣子', '香辣', '酸辣', '藤椒'],
        '酸甜': ['糖醋', '酸甜', '番茄', '咕咾', '菠萝', '柠檬'],
        '酱香': ['酱香', '酱', '京酱', '南乳', '腐乳', '照烧', '叉烧'],
        '孜然': ['孜然'],
        '蒜香': ['蒜香', '蒜蓉', '蒜泥', '金银蒜'],
        '咖喱': ['咖喱'],
    }
    for tag, kws in kw_map.items():
        if any(kw in title for kw in kws):
            tags.add(tag)
    return list(tags)[:10]


def match_table_to_recipe(table_data, recipe_title):
    """判断表格是否属于指定菜谱"""
    all_text = ' '.join([' '.join(row) for row in table_data])
    # 检查表格中是否包含菜谱标题的关键词
    title_keywords = recipe_title.replace('（', '').replace('）', '').replace('(', '').replace(')', '')
    # 取标题的前几个字作为关键词
    for kw_len in [4, 3, 2]:
        kw = title_keywords[:kw_len]
        if kw in all_text:
            return True
    return False


def parse_docx(filepath):
    doc = docx.Document(filepath)
    body = list(doc.element.body)

    # ── 第一遍：收集所有独立的表格 ──
    all_tables = []
    for i, el in enumerate(body):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'tbl':
            # 找到对应的 doc.tables 索引
            tidx = len(all_tables)
            if tidx < len(doc.tables):
                td = extract_table_data(doc.tables[tidx])
                all_tables.append({'body_index': i, 'table_index': tidx, 'data': td})

    print(f"   找到 {len(all_tables)} 个表格")

    # ── 第二遍：解析菜谱段落内容 ──
    recipes = []
    current = None
    current_section = None
    steps_phase = None
    step_items = []
    table_counter = 0  # 跟踪当前处理到第几个表格

    def flush_phase():
        nonlocal steps_phase, step_items
        if steps_phase and step_items:
            current['steps'].append({'phase': steps_phase, 'items': step_items})
            step_items = []
            steps_phase = None

    def save_recipe():
        nonlocal current
        if not current:
            return
        flush_phase()
        for k in ['_nutrition', '_extra']:
            current.pop(k, None)
        current['subcategory'] = guess_subcategory(current['title'])
        current['category'] = '主食' if current['subcategory'] in ('面食', '米饭') else '家常菜'
        current['tags'] = extract_tags(current['title'], current.get('ingredients', []))
        current['difficulty'] = '简单' if len(current.get('ingredients', [])) <= 5 else '中等'
        if current.get('cookingTime', 0) <= 0:
            current['cookingTime'] = 30
        recipes.append(current)
        current = None

    def new_recipe(title):
        nonlocal current, current_section, steps_phase, step_items
        save_recipe()
        current = {
            'title': title,
            'category': '', 'subcategory': '',
            'tags': [], 'difficulty': '中等', 'cookingTime': 0,
            'servings': '2-3人份', 'ingredients': [], 'steps': [],
            'tips': [], 'nutrition': [], 'extra': [],
        }
        current_section = None
        steps_phase = None
        step_items = []

    para_idx = 0

    for el in body:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag

        if tag == 'p':
            if para_idx >= len(doc.paragraphs):
                para_idx += 1
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1
            text = clean_text(para.text)
            if not text:
                continue

            # 获取字号
            sz = None
            for run in para.runs:
                if run.font.size:
                    sz = run.font.size
                    break

            # 菜谱标题
            if sz and sz >= 190000 and text not in SKIP and len(text) >= 2:
                if not re.match(r'^[0-9\s\.、，。！？]+$', text):
                    new_recipe(text)
                    continue

            if not current:
                continue

            # 章节标题
            if sz and sz >= 177000:
                if text.startswith('一、') or text.startswith('一 ') or '食材清单' in text:
                    flush_phase()
                    current_section = 'ingredients'
                    sv = extract_servings(text)
                    if sv:
                        current['servings'] = sv
                elif text.startswith('二、') or text.startswith('二 ') or '详细步骤' in text or '步骤' in text[:6]:
                    flush_phase()
                    current_section = 'steps'
                elif text.startswith('三、') or text.startswith('三 ') or '成功关键' in text or '风味解析' in text or '关键' in text[:4]:
                    flush_phase()
                    current_section = 'tips'
                elif text.startswith('四、') or text.startswith('四 ') or '营养' in text[:6]:
                    flush_phase()
                    current_section = 'nutrition'
                continue

            # 跳过"点击图片可查看完整电子表格"
            if '点击图片' in text or '电子表格' in text:
                continue

            # 处理章节内容
            if current_section == 'steps':
                phase_match = re.match(r'^第[一二三四五六七八九十\d]+部分[：:]\s*(.+)', text)
                if phase_match:
                    flush_phase()
                    steps_phase = phase_match.group(1).strip()
                    continue
                cleaned = re.sub(r'^[（(]?\d+[）).、\s]*', '', text).strip()
                cleaned = re.sub(r'^[-\s•◦▪▸►●○◆◇]*\s*', '', text).strip()
                if cleaned and len(cleaned) > 5:
                    if not steps_phase:
                        steps_phase = '制作步骤'
                    step_items.append(cleaned)
                continue

            elif current_section == 'tips':
                cleaned = re.sub(r'^[-\s•◦▪▸►●○◆◇]*\s*', '', text).strip()
                if cleaned and len(cleaned) > 5:
                    current['tips'].append(cleaned)
                continue

            elif current_section == 'nutrition':
                current.setdefault('_nutrition', []).append(text)
                continue

            elif current_section == 'ingredients':
                # 尝试从文本中解析食材
                # 格式: "食材名：用量（备注）" 或 "食材名 用量"
                cleaned = re.sub(r'^[-\s•◦▪▸►●○◆◇]*\s*', '', text).strip()
                if cleaned and len(cleaned) > 3:
                    m = re.match(r'^(.+?)[：:]\s*(.+?)$', cleaned)
                    if m:
                        name = m.group(1).strip()
                        amount = m.group(2).strip()
                        note = ''
                        nm = re.match(r'^(.+?)[（(](.+?)[）)]$', amount)
                        if nm:
                            amount = nm.group(1).strip()
                            note = nm.group(2).strip()
                        if len(name) <= 15:
                            current['ingredients'].append({'name': name, 'amount': amount, 'note': note})
                continue

        elif tag == 'tbl':
            # 表格在文档末尾，尝试匹配到当前菜谱
            if current and table_counter < len(all_tables):
                td = all_tables[table_counter]['data']
                table_counter += 1
                if not td or len(td) < 2:
                    continue

                hdr = td[0]
                hdr_text = ' '.join(hdr).lower().replace('\u200b', '')

                # 判断是否是食材表格
                is_ingredient = any(k in hdr_text for k in ['食材', '材料', '主料', '用量', '分量'])
                is_comparison = any(k in hdr_text for k in ['对比', '区别', '差异', '版本', '方法', '风味', '特点', '问题', '解决方案', '风格', '类型'])
                is_tip = any(k in hdr_text for k in ['关键', '技巧', '要点', '步骤', '阶段', '大阶段', '核心'])

                # 尝试匹配到当前菜谱
                matched = match_table_to_recipe(td, current['title'])

                if is_ingredient and not is_comparison and not is_tip:
                    ings = parse_ingredients(td)
                    if ings:
                        current['ingredients'].extend(ings)
                        # 从表格中估算烹饪时间
                        all_tbl_text = ' '.join([' '.join(row) for row in td])
                        t = extract_steps_time(all_tbl_text)
                        if t > current.get('cookingTime', 0):
                            current['cookingTime'] = t
                    continue

                if matched and is_tip:
                    for row in td[1:]:
                        for cell in row:
                            if cell and len(cell) > 5:
                                current['tips'].append(cell)
                    continue

                # 其他表格：如果匹配到菜谱，则作为补充内容
                if matched or is_comparison or is_tip:
                    current.setdefault('_extra', []).append({
                        'type': 'table',
                        'header': hdr,
                        'rows': td,
                    })
            else:
                # 超过 all_tables 范围的表格
                if table_counter < len(all_tables):
                    table_counter += 1

    # 保存最后一个菜谱
    save_recipe()

    # ── 第三遍：将未匹配的表格匹配到菜谱 ──
    # 重新遍历所有表格，为没有食材的菜谱匹配食材表格
    unmatched_tables = []
    for ti, t in enumerate(all_tables):
        td = t['data']
        if not td or len(td) < 2:
            continue
        hdr_text = ' '.join(td[0]).lower().replace('\u200b', '')
        is_ingredient = any(k in hdr_text for k in ['食材', '材料', '主料', '用量', '分量'])
        if is_ingredient:
            unmatched_tables.append((ti, td))

    for recipe in recipes:
        if recipe['ingredients']:
            continue  # 已有食材，跳过
        for ti, td in unmatched_tables:
            if match_table_to_recipe(td, recipe['title']):
                ings = parse_ingredients(td)
                if ings:
                    recipe['ingredients'] = ings
                    all_tbl_text = ' '.join([' '.join(row) for row in td])
                    t = extract_steps_time(all_tbl_text)
                    if t > recipe.get('cookingTime', 0):
                        recipe['cookingTime'] = t
                break

    # 尝试从步骤文本中提取烹饪时间
    for recipe in recipes:
        if recipe['cookingTime'] <= 30:
            all_steps = ' '.join([item for s in recipe['steps'] for item in s['items']])
            t = extract_steps_time(all_steps)
            if t > 0:
                recipe['cookingTime'] = t

    # 分配 ID
    for i, r in enumerate(recipes):
        r['id'] = i + 1

    return recipes


def main():
    parser = argparse.ArgumentParser(description='解析私房菜谱 .docx 为 JSON')
    parser.add_argument('docx_file', help='.docx 文件路径')
    parser.add_argument('-o', '--output', default='recipes.json', help='输出 JSON 路径')
    args = parser.parse_args()

    print(f'📖 读取: {args.docx_file}')
    recipes = parse_docx(args.docx_file)
    print(f'✅ 解析出 {len(recipes)} 道菜谱')

    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(recipes, f, ensure_ascii=False, indent=2)

    size_mb = os.path.getsize(args.output) / 1024 / 1024
    print(f'📄 输出: {args.output} ({size_mb:.1f} MB)')

    # 统计
    cats = {}
    with_ingredients = 0
    for r in recipes:
        if r['ingredients']:
            with_ingredients += 1
        k = f"{r['category']}/{r['subcategory']}"
        cats[k] = cats.get(k, 0) + 1
    for k, v in sorted(cats.items()):
        print(f'   {k}: {v} 道')
    print(f'\n📊 含食材清单: {with_ingredients}/{len(recipes)} 道')


if __name__ == '__main__':
    main()